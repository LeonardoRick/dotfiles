#!/usr/bin/env python3
"""
Sign a .docx file by inserting a signature image at the signature marker.

Usage:
    python sign_docx.py <docx_path> <signature_image_path> [--width cm] [--output path]

The script looks for common signature markers (e.g., "Atenciosamente", "Assinatura",
"____", "Signature") and places the image after the marker. Falls back to appending
at the end if no marker is found.

NEVER overwrites the original file — creates a _signed copy.
"""

import argparse
import os
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm


SIGNATURE_MARKERS = [
    "atenciosamente",
    "assinatura",
    "____",
    "signature",
    "signed by",
    "sincerely",
    "regards",
    "com os melhores cumprimentos",
    "cordialmente",
]


def find_signature_position(doc: Document) -> int | None:
    """Find the paragraph index containing a signature marker."""
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.lower().strip()
        for marker in SIGNATURE_MARKERS:
            if marker in text_lower:
                return i
    return None


def sign_docx(
    docx_path: str,
    signature_path: str,
    width_cm: float = 5.0,
    output_path: str | None = None,
) -> str:
    """Insert signature image into a .docx file.

    Args:
        docx_path: Path to the source .docx file
        signature_path: Path to the signature image (.png recommended)
        width_cm: Width of the signature image in centimeters
        output_path: Optional output path. Defaults to <original>_signed.docx

    Returns:
        The path to the signed document
    """
    docx_path = os.path.expanduser(docx_path)
    signature_path = os.path.expanduser(signature_path)

    if not os.path.exists(docx_path):
        print(f"Error: Document not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    if not os.path.exists(signature_path):
        print(f"Error: Signature image not found: {signature_path}", file=sys.stderr)
        sys.exit(1)

    if output_path is None:
        base = Path(docx_path)
        output_path = str(base.parent / f"{base.stem}_signed{base.suffix}")

    doc = Document(docx_path)
    marker_idx = find_signature_position(doc)

    if marker_idx is not None:
        # Check if there's a placeholder underline paragraph right after the marker
        next_idx = marker_idx + 1
        if next_idx < len(doc.paragraphs):
            next_text = doc.paragraphs[next_idx].text.strip()
            if all(c == "_" for c in next_text) and len(next_text) > 3:
                # Replace the underline placeholder with the signature
                target_para = doc.paragraphs[next_idx]
                target_para.clear()
            else:
                target_para = doc.paragraphs[next_idx]
        else:
            target_para = doc.add_paragraph()

        run = target_para.runs[0] if target_para.runs else target_para.add_run()
        run.add_picture(signature_path, width=Cm(width_cm))
    else:
        # No marker found — append at end
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(signature_path, width=Cm(width_cm))

    doc.save(output_path)
    print(f"Signed document saved to: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Sign a .docx document with an image signature")
    parser.add_argument("docx_path", help="Path to the .docx file to sign")
    parser.add_argument("signature_path", help="Path to the signature image (.png)")
    parser.add_argument("--width", type=float, default=5.0, help="Signature width in cm (default: 5.0)")
    parser.add_argument("--output", help="Output file path (default: <original>_signed.docx)")
    args = parser.parse_args()

    sign_docx(args.docx_path, args.signature_path, args.width, args.output)


if __name__ == "__main__":
    main()
